"""Build a clean Word delivery containing all 15 spoken-video transcripts."""

from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
SOURCE_DIR = ROOT / "data" / "transcripts" / "all15_whisper_fresh_20260630"
REPORT_PATH = ROOT / "reports" / "summary_context_ablation_15videos.json"
OUTPUT = ROOT / "reports" / "transcripts" / "beatriz_video_transcripts_all_15.docx"
PLAIN_TEXT_OUTPUT = ROOT / "reports" / "transcripts" / "beatriz_video_transcripts_all_15.txt"

VIDEO_IDS = [
    "KTNcYYHuBTY",
    "BZxZ_eEuJBM",
    "kpywdu1afas",
    "KeYfQ8Em9BU",
    "AqtgBbEI4wM",
    "G8GAsYkZlpE",
    "8mjcnxGMwFo",
    "LVLuqNH5iWw",
    "vwUV2IDLP8Q",
    "kv8pMEALdWo",
    "QqAPCIIdx-E",
    "QhB481jPAa4",
    "JTDa2uZS2gI",
    "TDHI-aieyfk",
    "uKfcS7-O6UE",
]

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 99, 110)


def set_run_font(run, *, size=11, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Calibri")
    rpr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_page_field(paragraph) -> None:
    paragraph.add_run("Page ")
    run = paragraph.add_run()
    for tag, value in (
        ("w:fldChar", ("w:fldCharType", "begin")),
        ("w:instrText", ("xml:space", "preserve")),
        ("w:fldChar", ("w:fldCharType", "separate")),
        ("w:t", None),
        ("w:fldChar", ("w:fldCharType", "end")),
    ):
        elem = OxmlElement(tag)
        if tag == "w:instrText":
            elem.set(qn(value[0]), value[1])
            elem.text = " PAGE "
        elif tag == "w:t":
            elem.text = "1"
        elif value:
            elem.set(qn(value[0]), value[1])
        run._r.append(elem)


def configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header_p = section.header.paragraphs[0]
    header_p.text = "MathE Research | Spoken Video Transcripts"
    header_p.paragraph_format.space_after = Pt(0)
    for run in header_p.runs:
        set_run_font(run, size=9, color=MUTED)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_after = Pt(0)
    add_page_field(footer_p)
    for run in footer_p.runs:
        set_run_font(run, size=9, color=MUTED)


def metadata_by_video() -> dict[str, dict]:
    report = json.loads(REPORT_PATH.read_text(encoding="utf-8"))
    first_combo = report["combos"][0]
    return {str(item["item_id"]): item for item in first_combo["items"]}


def add_metadata_line(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    label_run = p.add_run(f"{label}: ")
    set_run_font(label_run, size=10, color=DARK_BLUE, bold=True)
    value_run = p.add_run(value)
    set_run_font(value_run, size=10, color=MUTED)


def transcript_blocks(payload: dict, target_chars: int = 720) -> list[str]:
    segments = payload.get("segments") or []
    if not segments:
        text = str(payload.get("transcript") or "").strip()
        return [text] if text else []

    blocks: list[str] = []
    current: list[str] = []
    current_len = 0
    for segment in segments:
        text = str(segment.get("text") or "").strip()
        if not text:
            continue
        if current and current_len + len(text) + 1 > target_chars:
            blocks.append(" ".join(current))
            current = []
            current_len = 0
        current.append(text)
        current_len += len(text) + 1
    if current:
        blocks.append(" ".join(current))
    return blocks


def build() -> None:
    metadata = metadata_by_video()
    payloads = {}
    for video_id in VIDEO_IDS:
        path = SOURCE_DIR / f"{video_id}_whisper.json"
        payload = json.loads(path.read_text(encoding="utf-8"))
        if not str(payload.get("transcript") or "").strip():
            raise ValueError(f"Empty transcript: {video_id}")
        payloads[video_id] = payload

    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("MathE Video Transcripts")
    set_run_font(title_run, size=24, color=DARK_BLUE, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run("Complete spoken-audio transcripts for the 15-video evaluation set")
    set_run_font(subtitle_run, size=13, color=MUTED)

    add_metadata_line(doc, "Prepared for", "Beatriz")
    add_metadata_line(doc, "Prepared by", "Eren")
    add_metadata_line(doc, "Date", date.today().strftime("%d %B %Y"))
    add_metadata_line(doc, "Transcription method", "Local faster-whisper, distil-large-v3")

    intro = doc.add_paragraph()
    intro.paragraph_format.space_before = Pt(12)
    intro.paragraph_format.space_after = Pt(10)
    intro_run = intro.add_run(
        "This document contains the complete spoken transcripts for all 15 videos used in the "
        "keyword-evaluation spreadsheet. The audio was transcribed locally with a math-aware "
        "Whisper workflow and structured into paragraphs for readability. These texts are "
        "transcriptions of the spoken audio, not video summaries."
    )
    set_run_font(intro_run, size=11)

    plain_sections = []
    for index, video_id in enumerate(VIDEO_IDS, start=1):
        doc.add_page_break()
        item = metadata.get(video_id, {})
        payload = payloads[video_id]
        title_text = str(item.get("title") or video_id)
        url = str(item.get("url") or payload.get("url") or "")

        doc.add_heading(f"Video {index}: {title_text}", level=1)
        add_metadata_line(doc, "Video ID", video_id)
        add_metadata_line(doc, "URL", url)
        doc.add_heading("Transcript", level=2)

        blocks = transcript_blocks(payload)
        for text in blocks:
            p = doc.add_paragraph(text)
            p.paragraph_format.keep_together = False
            p.paragraph_format.widow_control = True

        plain_sections.append(
            f"VIDEO {index}: {title_text}\nVideo ID: {video_id}\nURL: {url}\n\n"
            + "\n\n".join(blocks)
        )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    PLAIN_TEXT_OUTPUT.write_text("\n\n" + ("\n\n" + "=" * 80 + "\n\n").join(plain_sections), encoding="utf-8")
    print(f"[OK] DOCX: {OUTPUT}")
    print(f"[OK] TXT:  {PLAIN_TEXT_OUTPUT}")


if __name__ == "__main__":
    build()
