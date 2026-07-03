"""Build a clean Word delivery containing three spoken-video transcripts."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
SOURCE_DIR = ROOT / "reports" / "transcripts" / "beatriz_pilot" / "gemini_primary"
OUTPUT = ROOT / "reports" / "transcripts" / "beatriz_video_transcripts_pilot.docx"

TRANSCRIPTS = [
    ("KTNcYYHuBTY", "Powers of the Imaginary Unit i"),
    ("LVLuqNH5iWw", "Local Extrema of a Multivariable Function"),
    ("uKfcS7-O6UE", "Quotient Rule for Differentiation"),
]

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 99, 110)


def set_run_font(run, *, size=11, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = BLUE
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = BLUE
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True


def add_page_field(paragraph) -> None:
    paragraph.add_run("Page ")
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "MathE Research | Spoken Video Transcripts"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    for run in hp.runs:
        set_run_font(run, size=9, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_after = Pt(0)
    add_page_field(fp)
    for run in fp.runs:
        set_run_font(run, size=9, color=MUTED)


def read_transcript(video_id: str) -> tuple[str, str, list[str]]:
    path = SOURCE_DIR / f"{video_id}_gemini_transcript.txt"
    lines = path.read_text(encoding="utf-8").splitlines()
    title = lines[0].removeprefix("# ").strip()
    url = next(line.split(": ", 1)[1] for line in lines if line.startswith("# url:"))
    body_start = next(i for i, line in enumerate(lines) if not line.startswith("#") and line.strip())
    body = "\n".join(lines[body_start:]).strip()
    paragraphs = [part.strip() for part in body.split("\n\n") if part.strip()]
    return title, url, paragraphs


def add_metadata_line(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    label_run = p.add_run(f"{label}: ")
    set_run_font(label_run, size=10, color=DARK_BLUE, bold=True)
    value_run = p.add_run(value)
    set_run_font(value_run, size=10, color=MUTED)


def build() -> None:
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("MathE Pilot Video Transcripts")
    set_run_font(title_run, size=24, color=DARK_BLUE, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run("Spoken audio transcribed into readable English text")
    set_run_font(subtitle_run, size=13, color=MUTED)

    add_metadata_line(doc, "Prepared for", "Beatriz")
    add_metadata_line(doc, "Prepared by", "Eren")
    add_metadata_line(doc, "Date", date.today().strftime("%d %B %Y"))

    intro = doc.add_paragraph()
    intro.paragraph_format.space_before = Pt(12)
    intro.paragraph_format.space_after = Pt(10)
    intro_run = intro.add_run(
        "This document contains the spoken transcripts of three pilot videos from the evaluated set. "
        "The audio was transcribed with a math-aware Gemini workflow and lightly structured into "
        "paragraphs for readability. The content below is a transcription, not a video summary."
    )
    set_run_font(intro_run, size=11)

    for index, (video_id, expected_title) in enumerate(TRANSCRIPTS, start=1):
        if index > 1:
            doc.add_page_break()
        title_text, url, paragraphs = read_transcript(video_id)
        section_title = title_text or expected_title
        doc.add_heading(f"Video {index}: {section_title}", level=1)
        add_metadata_line(doc, "Video ID", video_id)
        add_metadata_line(doc, "URL", url)

        transcript_heading = doc.add_heading("Transcript", level=2)
        transcript_heading.paragraph_format.space_before = Pt(14)
        for text in paragraphs:
            p = doc.add_paragraph(text)
            p.paragraph_format.keep_together = False
            p.paragraph_format.widow_control = True

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"[OK] {OUTPUT}")


if __name__ == "__main__":
    build()
